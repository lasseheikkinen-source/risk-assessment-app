import streamlit as st
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Värit (Rudus-brändi)
RUDUS_YELLOW = "#FFBB3C"
RUDUS_GRAY = "#E7E6E6"
CRH_BLUE_LIGHT = "#2860A8"
CRH_BLUE_DARK = "#041E42"

# Riskien hallinta hierarkia
CONTROL_HIERARCHY = [
    "🚫 Eliminointi - Poista vaara kokonaan",
    "🔄 Korvaus - Korvaa vaarallinen turvallisemmalla",
    "🛡️ Tekniset kontrollit - Eristä ihminen vaarasta",
    "📋 Hallinnolliset kontrollit - Muuta työtapaa",
    "🧤 Henkilösuojaimet - Suojavarusteet"
]

# Riskien data
RISKS_DATA = {
    "Painovoima": [
        {"riski": "Putoaminen korkealta", "todennäköisyys": 2, "vakavuus": 3, "riskin_arvo": 5},
        {"riski": "Esineen putoaminen pään päälle", "todennäköisyys": 2, "vakavuus": 3, "riskin_arvo": 5},
        {"riski": "Liukastuminen ja kaatuminen", "todennäköisyys": 3, "vakavuus": 2, "riskin_arvo": 5},
        {"riski": "Kuorman putoaminen", "todennäköisyys": 2, "vakavuus": 3, "riskin_arvo": 5},
    ],
    "Liike": [
        {"riski": "Jääminen liikkuvan koneen väliin", "todennäköisyys": 2, "vakavuus": 3, "riskin_arvo": 5},
        {"riski": "Osuminen liikkuvasta osasta", "todennäköisyys": 2, "vakavuus": 2, "riskin_arvo": 4},
        {"riski": "Pyörimisvaara (hiukset, vaatteet)", "todennäköisyys": 2, "vakavuus": 3, "riskin_arvo": 5},
        {"riski": "Nopean liikkeen aiheuttama vamma", "todennäköisyys": 2, "vakavuus": 2, "riskin_arvo": 4},
    ],
    "Mekaaninen": [
        {"riski": "Terävällä esineellä leikkautuminen", "todennäköisyys": 2, "vakavuus": 2, "riskin_arvo": 4},
        {"riski": "Puristuminen kahden pinnan väliin", "todennäköisyys": 2, "vakavuus": 3, "riskin_arvo": 5},
        {"riski": "Iskuvaara", "todennäköisyys": 2, "vakavuus": 2, "riskin_arvo": 4},
        {"riski": "Viilto tai naarmu", "todennäköisyys": 3, "vakavuus": 1, "riskin_arvo": 3},
    ],
    "Sähkö": [
        {"riski": "Sähköisku", "todennäköisyys": 1, "vakavuus": 3, "riskin_arvo": 3},
        {"riski": "Palovamma sähköstä", "todennäköisyys": 1, "vakavuus": 3, "riskin_arvo": 3},
        {"riski": "Palo sähköviasta", "todennäköisyys": 1, "vakavuus": 3, "riskin_arvo": 3},
        {"riski": "Kaatuminen sähköiskun seurauksena", "todennäköisyys": 1, "vakavuus": 2, "riskin_arvo": 2},
    ],
    "Paine": [
        {"riski": "Paineistetun säiliön räjähtäminen", "todennäköisyys": 1, "vakavuus": 3, "riskin_arvo": 3},
        {"riski": "Paineistetun nestesuihkun aiheuttama vamma", "todennäköisyys": 2, "vakavuus": 2, "riskin_arvo": 4},
        {"riski": "Paineilman aiheuttama vamma", "todennäköisyys": 2, "vakavuus": 2, "riskin_arvo": 4},
        {"riski": "Hydrauliikan vuoto ja palovamma", "todennäköisyys": 1, "vakavuus": 2, "riskin_arvo": 2},
    ],
    "Ääni": [
        {"riski": "Kuulovauriot pitkäaikaisesta melusta", "todennäköisyys": 3, "vakavuus": 2, "riskin_arvo": 5},
        {"riski": "Väsymys ja keskittymiskyvyn heikkeneminen", "todennäköisyys": 3, "vakavuus": 1, "riskin_arvo": 3},
        {"riski": "Kommunikaation vaikeutuminen", "todennäköisyys": 3, "vakavuus": 1, "riskin_arvo": 3},
        {"riski": "Stressin lisääntyminen", "todennäköisyys": 2, "vakavuus": 1, "riskin_arvo": 2},
    ],
    "Lämpötila": [
        {"riski": "Palovamma kuumasta pinnasta", "todennäköisyys": 2, "vakavuus": 2, "riskin_arvo": 4},
        {"riski": "Paleltuma kylmästä", "todennäköisyys": 1, "vakavuus": 2, "riskin_arvo": 2},
        {"riski": "Lämpöväsymys tai lämpöhalvaus", "todennäköisyys": 2, "vakavuus": 2, "riskin_arvo": 4},
        {"riski": "Hypotermia kylmässä ympäristössä", "todennäköisyys": 1, "vakavuus": 3, "riskin_arvo": 3},
    ],
    "Kemiallinen": [
        {"riski": "Kemikaalin aiheuttama palovamma", "todennäköisyys": 1, "vakavuus": 3, "riskin_arvo": 3},
        {"riski": "Myrkyllisen kaasun hengittäminen", "todennäköisyys": 1, "vakavuus": 3, "riskin_arvo": 3},
        {"riski": "Ihon tai silmien ärsytys", "todennäköisyys": 2, "vakavuus": 2, "riskin_arvo": 4},
        {"riski": "Kemiallinen reaktio ja räjähdys", "todennäköisyys": 1, "vakavuus": 3, "riskin_arvo": 3},
    ],
}

def get_risk_color(risk_value):
    if risk_value <= 2:
        return "#90EE90"
    elif risk_value <= 4:
        return "#FFD700"
    else:
        return "#FF6B6B"

def get_risk_level_text(risk_value):
    if risk_value <= 2:
        return "🟢 Pieni"
    elif risk_value <= 4:
        return "🟡 Keskisuuri"
    else:
        return "🔴 Suuri"

def create_word_report(selected_risks, comments):
    doc = Document()
    title = doc.add_heading("Riskien Arviointiraportti", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para = doc.add_paragraph(f"Päivämäärä: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    doc.add_heading("Tunnistetut Riskit", level=1)
    
    for risk_id, risk_data in selected_risks.items():
        doc.add_heading(f"⚠️ {risk_data['riski']}", level=2)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_table.rows[0].cells[0].text = "Energialähde"
        info_table.rows[0].cells[1].text = risk_data['energia']
        info_table.rows[1].cells[0].text = "Todennäköisyys"
        info_table.rows[1].cells[1].text = f"{risk_data['todennäköisyys']}/3"
        info_table.rows[2].cells[0].text = "Vakavuus"
        info_table.rows[2].cells[1].text = f"{risk_data['vakavuus']}/3"
        info_table.rows[3].cells[0].text = "Riskin arvo"
        info_table.rows[3].cells[1].text = f"{risk_data['riskin_arvo']}/6 - {get_risk_level_text(risk_data['riskin_arvo'])}"
        
        if risk_id in comments and comments[risk_id].get('kommentti'):
            doc.add_heading("Kommentit", level=3)
            doc.add_paragraph(comments[risk_id]['kommentti'])
        
        if risk_id in comments and comments[risk_id].get('toimenpiteet'):
            doc.add_heading("Toimenpiteet", level=3)
            for toimenpide in comments[risk_id]['toimenpiteet']:
                doc.add_paragraph(toimenpide, style='List Bullet')
        
        doc.add_paragraph()
    
    doc.add_heading("Riskien Hallinta Hierarkia", level=1)
    for i, level in enumerate(CONTROL_HIERARCHY, 1):
        doc.add_paragraph(level, style='List Number')
    
    return doc

st.set_page_config(page_title="Riskien Arviointi", layout="wide", initial_sidebar_state="expanded")
st.markdown(f"""
    <style>
    .main {{
        background-color: #f5f5f5;
    }}
    </style>
""", unsafe_allow_html=True)

if 'selected_risks' not in st.session_state:
    st.session_state.selected_risks = {}
if 'comments' not in st.session_state:
    st.session_state.comments = {}
if 'custom_risks' not in st.session_state:
    st.session_state.custom_risks = []

st.title("⚡ Riskien Arviointi - Pikatilanteissa")
st.markdown("Rudus - Työturvallisuus")
st.markdown("---")

st.header("Vaihe 1: Valitse Energialähde")
col1, col2, col3, col4 = st.columns(4)
energies = list(RISKS_DATA.keys())

for i, energy in enumerate(energies):
    col = [col1, col2, col3, col4][i % 4]
    with col:
        if st.button(f"⚡ {energy}", key=f"energy_{energy}", use_container_width=True):
            st.session_state.current_energy = energy

st.markdown("---")

if 'current_energy' in st.session_state:
    energy = st.session_state.current_energy
    st.header(f"Vaihe 2: Valitse Riskit - {energy}")
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("Valmiit riskit:")
        for risk in RISKS_DATA[energy]:
            risk_id = f"{energy}_{risk['riski']}"
            col_check, col_info = st.columns([1, 3])
            
            with col_check:
                is_selected = st.checkbox(
                    f"✓",
                    value=risk_id in st.session_state.selected_risks,
                    key=f"check_{risk_id}"
                )
            
            with col_info:
                risk_color = get_risk_color(risk['riskin_arvo'])
                st.markdown(f"""
                    <div style="background-color: {risk_color}; padding: 10px; border-radius: 5px; color: black;">
                    <b>{risk['riski']}</b><br>
                    Todennäköisyys: {risk['todennäköisyys']}/3 | Vakavuus: {risk['vakavuus']}/3 | Arvo: {risk['riskin_arvo']}/6
                    </div>
                """, unsafe_allow_html=True)
            
            if is_selected and risk_id not in st.session_state.selected_risks:
                st.session_state.selected_risks[risk_id] = {
                    'energia': energy,
                    'riski': risk['riski'],
                    'todennäköisyys': risk['todennäköisyys'],
                    'vakavuus': risk['vakavuus'],
                    'riskin_arvo': risk['riskin_arvo']
                }
            elif not is_selected and risk_id in st.session_state.selected_risks:
                del st.session_state.selected_risks[risk_id]
    
    with col2:
        st.subheader("Oma riski:")
        custom_risk = st.text_input("Kirjoita oma riski:", key="custom_risk_input")
        custom_prob = st.slider("Todennäköisyys:", 1, 3, 2, key="custom_prob")
        custom_sev = st.slider("Vakavuus:", 1, 3, 2, key="custom_sev")
        
        if st.button("➕ Lisää oma riski"):
            if custom_risk:
                risk_id = f"custom_{len(st.session_state.custom_risks)}"
                risk_value = custom_prob * custom_sev
                if risk_value > 6:
                    risk_value = 6
                
                st.session_state.selected_risks[risk_id] = {
                    'energia': energy,
                    'riski': custom_risk,
                    'todennäköisyys': custom_prob,
                    'vakavuus': custom_sev,
                    'riskin_arvo': risk_value
                }
                st.session_state.custom_risks.append(custom_risk)
                st.success(f"✅ Riski lisätty! Arvo: {risk_value}/6")

st.markdown("---")

if st.session_state.selected_risks:
    st.header("Vaihe 3: Kommentoi Valittuja Riskejä")
    
    for risk_id, risk_data in st.session_state.selected_risks.items():
        with st.expander(f"⚠️ {risk_data['riski']} ({risk_data['energia']}) - Arvo: {risk_data['riskin_arvo']}/6"):
            comment = st.text_area(
                "Kommentti:",
                value=st.session_state.comments.get(risk_id, {}).get('kommentti', ''),
                key=f"comment_{risk_id}",
                height=100
            )
            
            st.subheader("Toimenpiteet (Riskien Hallinta Hierarkia):")
            actions = st.session_state.comments.get(risk_id, {}).get('toimenpiteet', [])
            
            for i, hierarchy_level in enumerate(CONTROL_HIERARCHY):
                action = st.text_input(
                    hierarchy_level,
                    value=actions[i] if i < len(actions) else '',
                    key=f"action_{risk_id}_{i}"
                )
                if action:
                    if i >= len(actions):
                        actions.append(action)
                    else:
                        actions[i] = action
            
            if risk_id not in st.session_state.comments:
                st.session_state.comments[risk_id] = {}
            
            st.session_state.comments[risk_id]['kommentti'] = comment
            st.session_state.comments[risk_id]['toimenpiteet'] = [a for a in actions if a]
    
    st.markdown("---")
    st.header("Vaihe 4: Vie Raportti")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📄 Lataa Word-raportti", use_container_width=True):
            doc = create_word_report(st.session_state.selected_risks, st.session_state.comments)
            doc.save("Riskien_Arviointi.docx")
            
            with open("Riskien_Arviointi.docx", "rb") as f:
                st.download_button(
                    label="⬇️ Lataa Word-tiedosto",
                    data=f.read(),
                    file_name="Riskien_Arviointi.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    with col2:
        if st.button("🔄 Tyhjennä kaikki", use_container_width=True):
            st.session_state.selected_risks = {}
            st.session_state.comments = {}
            st.session_state.custom_risks = []
            st.rerun()

else:
    st.info("👈 Valitse energialähde ja riskit aloittaaksesi")
